from datetime import date, timedelta
from django.shortcuts import render, redirect
from django.http import HttpResponse
from docx import Document

from .models import Factura, ProductoFactura, SemanaHistorial


# ============================================================
# CREAR FACTURA
# ============================================================
def crear_factura(request):
    if request.method == 'POST':
        fecha = request.POST.get('fecha')
        proveedor = request.POST.get('proveedor')
        numero_factura = request.POST.get('numero_factura')
        observaciones = request.POST.get('observaciones', '')

        nombres = request.POST.getlist('producto_nombre[]')
        precios = request.POST.getlist('producto_precio[]')

        # Validar datos
        if fecha and proveedor and numero_factura and any(n.strip() for n in nombres):
            factura = Factura.objects.create(
                fecha=fecha,
                proveedor=proveedor,
                numero_factura=numero_factura,
                observaciones=observaciones
            )

            # Guardar productos
            for nombre, precio in zip(nombres, precios):
                nombre = nombre.strip()
                if not nombre or not precio:
                    continue

                ProductoFactura.objects.create(
                    factura=factura,
                    nombre_producto=nombre,
                    precio=precio
                )

            return redirect('crear_factura')  # recarga limpia

    return render(request, 'core/crear_factura.html')



# ============================================================
# RESUMEN SEMANAL (Semana actual + Semana anterior)
# ============================================================
def resumen_semanal(request):

    hoy = date.today()

    # -------------------------
    # SEMANA ACTUAL
    # -------------------------
    inicio_semana = hoy - timedelta(days=hoy.weekday())
    fin_semana = inicio_semana + timedelta(days=6)

    facturas_semana = Factura.objects.filter(
        fecha__range=[inicio_semana, fin_semana]
    ).prefetch_related('productos')

    total_actual = sum(f.total for f in facturas_semana)

    # -------------------------
    # SEMANA ANTERIOR
    # -------------------------
    inicio_semana_anterior = inicio_semana - timedelta(days=7)
    fin_semana_anterior = fin_semana - timedelta(days=7)

    facturas_semana_anterior = Factura.objects.filter(
        fecha__range=[inicio_semana_anterior, fin_semana_anterior]
    ).prefetch_related('productos')

    total_anterior = sum(f.total for f in facturas_semana_anterior)

    # -------------------------
    # GUARDAR SEMANA PASADA EN HISTORIAL (solo lunes)
    # -------------------------
    if hoy.weekday() == 0:  # lunes
        semana_existente = SemanaHistorial.objects.filter(
            inicio_semana=inicio_semana_anterior
        ).exists()

        if not semana_existente:
            SemanaHistorial.objects.create(
                inicio_semana=inicio_semana_anterior,
                fin_semana=fin_semana_anterior,
                total_semana=total_anterior
            )

    return render(request, 'core/resumen_semanal.html', {
        # Semana actual
        'facturas_semana': facturas_semana,
        'total_actual': total_actual,
        'inicio_semana': inicio_semana,
        'fin_semana': fin_semana,

        # Semana anterior
        'facturas_semana_anterior': facturas_semana_anterior,
        'total_anterior': total_anterior,
        'inicio_semana_anterior': inicio_semana_anterior,
        'fin_semana_anterior': fin_semana_anterior,
    })



# ============================================================
# HISTORIAL COMPLETO
# ============================================================
def historial_semanas(request):
    historial = SemanaHistorial.objects.all().order_by('-inicio_semana')

    return render(request, 'core/historial_semanas.html', {
        'historial': historial
    })



# ============================================================
# EXPORTAR SEMANA ACTUAL A WORD (Solo títulos + total)
# ============================================================
def exportar_semana_word(request):
    hoy = date.today()

    # Semana actual
    inicio_semana = hoy - timedelta(days=hoy.weekday())
    fin_semana = inicio_semana + timedelta(days=6)

    facturas = Factura.objects.filter(
        fecha__range=[inicio_semana, fin_semana]
    ).prefetch_related('productos')

    # Crear documento Word
    doc = Document()
    doc.add_heading('Resumen semanal de facturas', level=1)

    doc.add_paragraph(f"Semana del {inicio_semana} al {fin_semana}\n")

    total_semana = 0

    for factura in facturas:
        total_semana += factura.total

        titulo = (
            f"{factura.fecha.strftime('%d/%m/%Y')} — "
            f"Factura {factura.numero_factura} — {factura.proveedor}"
        )

        doc.add_paragraph(titulo)
        doc.add_paragraph(f"Total: $ {int(factura.total)}")
        doc.add_paragraph("-----------------------------------")

    doc.add_paragraph("")
    doc.add_heading(f"TOTAL SEMANAL: $ {int(total_semana)}", level=2)

    # Respuesta HTTP
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=Resumen_Semanal.docx'
    doc.save(response)

    return response






def editar_factura(request, factura_id):
    factura = Factura.objects.get(id=factura_id)

    if request.method == 'POST':
        factura.fecha = request.POST.get('fecha')
        factura.proveedor = request.POST.get('proveedor')
        factura.numero_factura = request.POST.get('numero_factura')
        factura.observaciones = request.POST.get('observaciones', '')
        factura.save()

        # BORRAR todos los productos antiguos
        factura.productos.all().delete()

        # CREAR los productos nuevos (editados)
        nombres = request.POST.getlist('producto_nombre[]')
        precios = request.POST.getlist('producto_precio[]')

        for nombre, precio in zip(nombres, precios):
            nombre = nombre.strip()
            if not nombre or not precio:
                continue

            ProductoFactura.objects.create(
                factura=factura,
                nombre_producto=nombre,
                precio=precio
            )

        return redirect('resumen_semanal')

    return render(request, 'core/editar_factura.html', {
        'factura': factura
    })



